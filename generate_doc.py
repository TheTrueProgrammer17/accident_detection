from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_doc():
    doc = Document()
    
    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Arial'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 51, 102)

    # 1. Cover Page
    doc.add_paragraph('\n\n\n\n\n\n\n\n')
    title = doc.add_paragraph('Project Name: SadakSuraksha Grid – AI Accident Detection Prototype')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True
    
    team = doc.add_paragraph('\nTeam Name: [Your Team Name]')
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team.runs[0].font.size = Pt(14)

    hackathon = doc.add_paragraph('Hackathon Name: [Hackathon Name]')
    hackathon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hackathon.runs[0].font.size = Pt(14)
    
    date = doc.add_paragraph('Submission Date: [Date]')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.runs[0].font.size = Pt(14)
    
    doc.add_page_break()

    # 2. Project Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph('The prototype demonstrates how AI-powered smart infrastructure can automatically detect road accidents and coordinate emergency response using distributed communication nodes and computer vision.')

    # 3. Problem Statement
    doc.add_heading('2. Problem Statement', level=1)
    doc.add_paragraph('Road accidents often remain unnoticed for crucial minutes due to dependence on manual reporting systems, leading to delayed emergency response and preventable fatalities.')

    # 4. Current Prototype Implementation
    doc.add_heading('3. Current Prototype Implementation', level=1)
    doc.add_paragraph('The current prototype is a software-based implementation that showcases the core functionality of the system. It operates as follows:')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Uses recorded traffic/road footage as input.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('AI model processes video frames in real time.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Accident events are detected using computer vision.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('3 MQTT nodes simulate distributed smart streetlights.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Nodes communicate with each other through MQTT messaging.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Alerts are transmitted between nodes upon accident detection.')

    # 5. System Workflow
    doc.add_heading('4. System Workflow', level=1)
    doc.add_paragraph('Video Feed → AI Detection Module → MQTT Node Communication → Alert Generation → Dashboard/Authority Notification')

    # 6. Technical Architecture
    doc.add_heading('5. Technical Architecture', level=1)
    doc.add_paragraph('The current software architecture follows a distributed data flow pipeline:')
    doc.add_paragraph('Video Input → OpenCV/YOLO Processing → MQTT Broker → Node Communication → Alert System\n')
    
    doc.add_paragraph('Node Roles:', style='Heading 2')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Node 1:').bold = True
    p.add_run(' Detection node (Runs AI models on video feeds)')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Node 2:').bold = True
    p.add_run(' Communication/relay node (Simulates mesh routing)')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Node 3:').bold = True
    p.add_run(' Alert/dashboard node (Receives signals and displays alerts)')

    # 7. Technologies Used
    doc.add_heading('6. Technologies Used', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technology'
    hdr_cells[1].text = 'Purpose'
    # add rows
    techs = [
        ('Python', 'Core programming language for AI and node logic'),
        ('OpenCV', 'Video frame extraction and image processing'),
        ('YOLO', 'Object detection and accident scenario analysis'),
        ('MQTT', 'Lightweight messaging protocol for node-to-node communication'),
        ('Node.js', 'Backend services and API handling (if applicable)'),
        ('React.js', 'Frontend dashboard UI (if applicable)'),
        ('Firebase', 'Real-time database and user authentication (if applicable)'),
        ('WebSockets', 'Real-time dashboard updates'),
        ('REST APIs', 'System integration and data retrieval')
    ]
    for tech, purpose in techs:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = purpose

    doc.add_paragraph('\n')

    # 8. Accident Detection Logic
    doc.add_heading('7. Accident Detection Logic', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Frame-by-frame analysis').bold = True
    p.add_run(': Continuous processing of video input to capture rapid events.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Object detection').bold = True
    p.add_run(': Identification of vehicles, pedestrians, and road boundaries.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Sudden motion/collision analysis').bold = True
    p.add_run(': Tracking velocity vectors and sudden trajectory changes indicating impact.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('AI-based abnormal event detection').bold = True
    p.add_run(': Recognizing patterns associated with accidents rather than normal traffic flow.')

    # 9. MQTT Communication System
    doc.add_heading('8. MQTT Communication System', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Publish/Subscribe mechanism allows efficient, decoupled messaging.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Ensures real-time message transfer with minimal latency.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Enables distributed node communication, mimicking physical smart streetlights.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Event propagation occurs seamlessly between nodes upon critical incident detection.')

    # 10. Alert Workflow
    doc.add_heading('9. Alert Workflow', level=1)
    p = doc.add_paragraph(style='List Number')
    p.add_run('Accident detected by Node 1.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('MQTT message (alert payload) published to the broker.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Other nodes (Node 2, Node 3) receive the alert via subscription.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Emergency warning generated on the dashboard system.')
    p = doc.add_paragraph(style='List Number')
    p.add_run('Future scope: Direct integration with local emergency service APIs.')

    # 11. Future Hardware Vision
    doc.add_heading('10. Future Hardware Vision', level=1)
    doc.add_paragraph('While the current implementation is a software-based prototype, the final product is designed to be deployed as physical infrastructure. The future hardware vision includes:')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Deployment of Smart Streetlights equipped with sensors and computing modules.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Integration of Edge AI cameras for on-site video processing without heavy cloud reliance.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Utilization of ESP32-S3 microcontrollers as physical nodes.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Addition of GSM/GPS modules for independent cellular connectivity and precise location tagging.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Incorporation of physical warning lights (strobes/LEDs) to alert oncoming traffic.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Solar-powered roadside units for sustainable operation.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Implementation of a physical mesh communication network using LoRa or Zigbee for redundancy.')

    # 12. Implementation Code Snippets
    doc.add_heading('11. Implementation Code Snippets', level=1)
    
    doc.add_paragraph('The following snippet demonstrates how an alert payload is structured before transmission over MQTT:', style='Normal')
    
    code1 = doc.add_paragraph('''def build_alert() -> dict:
    """Construct the standard JSON alert payload."""
    return {
        "node":      NODE_NAME,
        "event":     "accident_detected",
        "location":  LOCATION,
        "severity":  SEVERITY,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }''')
    code1.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph('\nThe snippet below illustrates the frame-by-frame collision detection logic and alert triggering:', style='Normal')
    
    code2 = doc.add_paragraph('''if detector.is_collision(has_vehicles, motion_intensity):
    accident_frame_counter += 1
else:
    accident_frame_counter = 0

if accident_frame_counter >= COLLISION_FRAME_THRESHOLD:
    collision_detected = True

if collision_detected and alert_sent == False:
    alert = build_alert()
    publisher.publish_alert(alert)
    alert_sent = True''')
    code2.runs[0].font.name = 'Courier New'

    # 13. Feasibility Analysis
    doc.add_heading('12. Feasibility Analysis', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Existing CCTV infrastructure can be reused, reducing initial deployment costs.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Modular smart node deployment allows for gradual rollout rather than total infrastructure overhaul.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('MQTT protocol guarantees scalable communication, capable of handling thousands of nodes.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('The system can initially be piloted on high-risk highways and major intersections.')

    # 14. Limitations of Current Prototype
    doc.add_heading('13. Limitations of Current Prototype', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Currently utilizes pre-recorded video feeds instead of live IP camera streams.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Hardware nodes and network latency are simulated via software.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Integration with actual emergency services (police, ambulance dispatch) is not yet live.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('AI accuracy is contingent on the quality and diversity of the training data used for the prototype models.')

    # 15. Future Scope
    doc.add_heading('14. Future Scope', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Real-time CCTV integration using RTSP streams.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Edge AI processing directly on microcontroller units (e.g., Coral TPU, Jetson Nano).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Automatic ambulance dispatch system integration.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Advanced traffic analytics (flow tracking, congestion prediction).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Vehicle identification system (license plate recognition for involved vehicles).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('ArogyaVault integration for rapid medical data retrieval of accident victims.')

    # 16. Conclusion
    doc.add_heading('15. Conclusion', level=1)
    doc.add_paragraph('The SadakSuraksha Grid prototype successfully demonstrates the feasibility of an AI-powered distributed accident detection system. By leveraging computer vision and MQTT communication, this software prototype establishes a robust foundation for a future hardware-based smart infrastructure ecosystem. It proves that real-time event detection and automated alert propagation can significantly bridge the gap between incident occurrence and emergency response, ultimately aiming to save lives on the road.')

    doc.save('SadakSuraksha_Grid_Technical_Documentation.docx')

if __name__ == '__main__':
    create_doc()
